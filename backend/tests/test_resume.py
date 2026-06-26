"""Tests for the resume parsing and scoring engine."""

from __future__ import annotations

import io
import json
from unittest.mock import AsyncMock, patch

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pytest

from backend.core.exceptions import FileValidationError, ResumeParsingError
from backend.models.schemas import CandidateTier, ExperienceEntry, ParsedResume
from backend.resume.parser import (
    _extract_certifications,
    _extract_docx_text,
    _extract_education,
    _extract_email,
    _extract_experience,
    _extract_name,
    _extract_phone,
    _extract_profile_urls,
    _extract_uri_values,
    _is_supported_profile_url,
    _extract_skills,
    parse_resume,
)
from backend.resume.scorer import (
    _assign_tier,
    _compute_keyword_match,
    _cosine_similarity,
    _score_education_certs,
    _score_experience,
    _score_format_completeness,
    score_resume,
)


# ---------------------------------------------------------------------------
# Parser Tests
# ---------------------------------------------------------------------------

class TestParserHelpers:
    """Test individual parser helper functions."""

    def test_extract_email(self) -> None:
        """Test email extraction from text."""
        text = "Contact me at john.doe@example.com or backup@site.org"
        assert _extract_email(text) == "john.doe@example.com"

    def test_extract_email_none(self) -> None:
        """Test email extraction when no email present."""
        assert _extract_email("No email here") == ""

    def test_extract_phone(self) -> None:
        """Test phone extraction from text."""
        text = "Call me at 555-123-4567 or (555) 987-6543"
        phone = _extract_phone(text)
        assert "555" in phone
        assert "123" in phone
        assert "4567" in phone

    def test_extract_profile_urls(self) -> None:
        """Test public profile URL extraction from resume text."""
        text = """
        GitHub: github.com/janesmith
        LinkedIn: https://linkedin.com/in/janesmith
        ML: huggingface.co/janesmith
        Practice: leetcode.com/u/janesmith.
        """
        urls = _extract_profile_urls(text)
        assert "https://github.com/janesmith" in urls
        assert "https://linkedin.com/in/janesmith" in urls
        assert "https://huggingface.co/janesmith" in urls
        assert "https://leetcode.com/u/janesmith" in urls

    def test_extract_profile_urls_skips_empty_profile_roots(self) -> None:
        """Test empty profile placeholders are ignored."""
        text = "GitHub: https://github.com/ LinkedIn: https://linkedin.com/in/ Profile: github.com/janesmith"
        urls = _extract_profile_urls(text)
        assert "https://github.com/janesmith" in urls
        assert "https://github.com/" not in urls
        assert "https://linkedin.com/in/" not in urls

    def test_is_supported_profile_url_accepts_owned_work(self) -> None:
        """Test owned repository/project URLs are still accepted as useful evidence."""
        assert _is_supported_profile_url("https://github.com/janesmith/project")
        assert _is_supported_profile_url("https://huggingface.co/spaces/janesmith/demo")
        assert not _is_supported_profile_url("https://github.com/topics/python")

    def test_extract_uri_values_from_nested_link_metadata(self) -> None:
        """Test URL extraction from nested PDF annotation metadata."""
        metadata = {
            "data": {
                "A": {
                    "URI": "https://www.linkedin.com/in/janesmith",
                },
            },
            "extra": ["https://github.com/janesmith"],
        }
        urls = _extract_uri_values(metadata)
        assert "https://www.linkedin.com/in/janesmith" in urls
        assert "https://github.com/janesmith" in urls

    def test_extract_docx_text_includes_hyperlink_targets(self) -> None:
        """Test DOCX hyperlink relationships are included even when only label text is visible."""
        doc = Document()
        doc.add_paragraph("Jane Smith")
        doc.add_paragraph("LinkedIn GitHub")
        doc.part.relate_to("https://www.linkedin.com/in/janesmith", RT.HYPERLINK, is_external=True)
        doc.part.relate_to("https://github.com/janesmith", RT.HYPERLINK, is_external=True)

        content = io.BytesIO()
        doc.save(content)

        text = _extract_docx_text(content.getvalue())
        urls = _extract_profile_urls(text)
        assert "https://www.linkedin.com/in/janesmith" in urls
        assert "https://github.com/janesmith" in urls

    def test_extract_name(self) -> None:
        """Test name extraction from resume."""
        text = "Jane Smith\nEmail: jane@example.com\n\nSummary\nExperienced developer."
        name = _extract_name(text)
        assert name == "Jane Smith"

    def test_extract_skills(self, sample_resume_text: str) -> None:
        """Test skill extraction from resume."""
        skills = _extract_skills(sample_resume_text)
        assert "python" in skills
        assert "docker" in skills
        assert "kubernetes" in skills
        assert "fastapi" in skills or "django" in skills

    def test_extract_experience(self) -> None:
        """Test experience extraction from resume."""
        text = """EXPERIENCE

Senior Backend Engineer | TechCorp Inc. | Jan 2021 - Present
Built scalable APIs using Python and FastAPI.

Backend Engineer | StartupXYZ | Jun 2018 - Dec 2020
Built REST APIs serving 500K+ users.

EDUCATION
BS Computer Science
"""
        exp = _extract_experience(text)
        assert len(exp) > 0
        assert any(e.company and "TechCorp" in e.company for e in exp)

    def test_extract_education(self, sample_resume_text: str) -> None:
        """Test education extraction from resume."""
        edu = _extract_education(sample_resume_text)
        assert len(edu) > 0
        assert any("Computer Science" in e.field for e in edu)

    def test_extract_certifications(self, sample_resume_text: str) -> None:
        """Test certification extraction from resume."""
        certs = _extract_certifications(sample_resume_text)
        assert len(certs) > 0
        assert any("AWS" in c.name for c in certs)

    def test_extract_name_no_resume(self) -> None:
        """Test name extraction with non-resume text."""
        text = "This is just some random text without a proper name header"
        name = _extract_name(text)
        # Should either find nothing or something reasonable
        assert isinstance(name, str)


class TestFileValidation:
    """Test file validation in parser."""

    def test_unsupported_extension(self) -> None:
        """Test that unsupported file extensions raise error."""
        from backend.resume.parser import _validate_file
        with pytest.raises(FileValidationError, match="Unsupported file format"):
            _validate_file("resume.txt", b"some content")

    def test_file_too_large(self) -> None:
        """Test that oversized files raise error."""
        from backend.resume.parser import _validate_file
        large_content = b"x" * (15 * 1024 * 1024)  # 15MB
        with pytest.raises(FileValidationError, match="too large"):
            _validate_file("resume.pdf", large_content)

    def test_valid_pdf(self) -> None:
        """Test that valid PDF passes validation."""
        from backend.resume.parser import _validate_file
        ext = _validate_file("resume.pdf", b"%PDF-1.4 minimal pdf content")
        assert ext == "pdf"

    def test_valid_docx(self) -> None:
        """Test that valid DOCX passes validation."""
        from backend.resume.parser import _validate_file
        ext = _validate_file("resume.docx", b"PK\x03\x04 docx content")
        assert ext == "docx"


# ---------------------------------------------------------------------------
# Scorer Tests
# ---------------------------------------------------------------------------

class TestKeywordMatch:
    """Test keyword matching logic."""

    def test_perfect_match(self) -> None:
        """Test perfect skill match gives full score."""
        resume_skills = ["python", "docker", "kubernetes"]
        job_skills = ["python", "docker", "kubernetes"]
        score = _compute_keyword_match(resume_skills, job_skills)
        assert score == 25.0

    def test_partial_match(self) -> None:
        """Test partial skill match gives proportional score."""
        resume_skills = ["python", "docker"]
        job_skills = ["python", "docker", "kubernetes", "aws"]
        score = _compute_keyword_match(resume_skills, job_skills)
        assert 10.0 <= score <= 13.0

    def test_no_job_skills(self) -> None:
        """Test that empty job skills list gives full score."""
        score = _compute_keyword_match(["python"], [])
        assert score == 25.0

    def test_no_match(self) -> None:
        """Test zero match gives zero score."""
        resume_skills = ["ruby"]
        job_skills = ["python", "java"]
        score = _compute_keyword_match(resume_skills, job_skills)
        assert score == 0.0


class TestCosineSimilarity:
    """Test cosine similarity computation."""

    def test_identical_vectors(self) -> None:
        """Test identical vectors have similarity 1.0."""
        vec = [1.0, 2.0, 3.0]
        sim = _cosine_similarity(vec, vec)
        assert sim == 1.0

    def test_orthogonal_vectors(self) -> None:
        """Test orthogonal vectors have similarity 0.0."""
        vec_a = [1.0, 0.0, 0.0]
        vec_b = [0.0, 1.0, 0.0]
        sim = _cosine_similarity(vec_a, vec_b)
        assert abs(sim) < 0.01

    def test_opposite_vectors(self) -> None:
        """Test opposite vectors have similarity -1.0."""
        vec_a = [1.0, 2.0, 3.0]
        vec_b = [-1.0, -2.0, -3.0]
        sim = _cosine_similarity(vec_a, vec_b)
        assert abs(sim - (-1.0)) < 0.01


class TestExperienceScoring:
    """Test experience depth scoring."""

    def test_senior_experience(self) -> None:
        """Test senior experience scoring."""
        exp = [
            ExperienceEntry(
                company="Corp", title="Senior Engineer",
                start_date="2020", end_date="Present",
                description="Led team of 10", years=4.0, is_senior=True,
            ),
        ]
        score = _score_experience(exp, "senior python engineer kubernetes")
        assert score >= 15.0

    def test_no_experience(self) -> None:
        """Test scoring with no experience."""
        score = _score_experience([], "python engineer")
        assert score == 0.0


class TestEducationCertScoring:
    """Test education and certification scoring."""

    def test_phd(self) -> None:
        """Test PhD gets maximum education score."""
        from backend.models.schemas import EducationEntry
        edu = [EducationEntry(degree="PhD", field="CS", institution="MIT", year="2020")]
        score = _score_education_certs(edu, [], "")
        assert score >= 10.0

    def test_bachelor_with_certs(self) -> None:
        """Test bachelor's with relevant certifications."""
        from backend.models.schemas import CertificationEntry, EducationEntry
        edu = [EducationEntry(degree="Bachelor of Science", field="CS", institution="Uni", year="2020")]
        certs = [CertificationEntry(name="AWS Certified Solutions Architect", issuer="AWS", year="2022")]
        score = _score_education_certs(edu, certs, "")
        assert score >= 8.0

    def test_no_education(self) -> None:
        """Test no education gives low score."""
        score = _score_education_certs([], [], "")
        assert score == 0.0


class TestFormatScoring:
    """Test format and completeness scoring."""

    def test_good_parse(self, parsed_resume: ParsedResume) -> None:
        """Test good parse quality gives high score."""
        score = _score_format_completeness(parsed_resume)
        assert score >= 7.0

    def test_poor_parse(self) -> None:
        """Test poor parse quality gives low score."""
        resume = ParsedResume(parse_quality="poor", sections_found=[], raw_text="")
        score = _score_format_completeness(resume)
        assert score <= 3.0


class TestTierAssignment:
    """Test tier assignment logic."""

    def test_tier_1(self) -> None:
        """Test score >= 90 gives Tier 1."""
        from backend.resume.scorer import _assign_tier
        assert _assign_tier(95.0) == CandidateTier.TIER_1
        assert _assign_tier(90.0) == CandidateTier.TIER_1

    def test_tier_2(self) -> None:
        """Test score 75-89 gives Tier 2."""
        from backend.resume.scorer import _assign_tier
        assert _assign_tier(80.0) == CandidateTier.TIER_2
        assert _assign_tier(75.0) == CandidateTier.TIER_2

    def test_tier_3(self) -> None:
        """Test score 60-74 gives Tier 3."""
        from backend.resume.scorer import _assign_tier
        assert _assign_tier(65.0) == CandidateTier.TIER_3
        assert _assign_tier(60.0) == CandidateTier.TIER_3

    def test_reject(self) -> None:
        """Test score < 60 gives Reject."""
        from backend.resume.scorer import _assign_tier
        assert _assign_tier(50.0) == CandidateTier.REJECT
        assert _assign_tier(0.0) == CandidateTier.REJECT


# ---------------------------------------------------------------------------
# Integration Tests
# ---------------------------------------------------------------------------

class TestResumeScoringIntegration:
    """Integration tests for the full scoring pipeline."""

    @pytest.mark.asyncio
    async def test_score_resume_with_embedding_mock(
        self,
        parsed_resume: ParsedResume,
        sample_job_description: str,
        mock_openai_embedding_response: dict,
    ) -> None:
        """Test full resume scoring with mocked embeddings."""
        with patch("backend.resume.scorer._get_embedding", new_callable=AsyncMock) as mock_embed:
            mock_embed.side_effect = [
                [0.1] * 1536,  # resume embedding
                [0.15] * 1536,  # job description embedding
            ]

            result = await score_resume(parsed_resume, sample_job_description)

            assert 0 <= result.total_score <= 100
            assert result.breakdown is not None
            assert result.tier is not None
            assert result.processing_time_ms >= 0
            assert result.cached is False

    @pytest.mark.asyncio
    async def test_score_resume_embedding_fallback(
        self,
        parsed_resume: ParsedResume,
        sample_job_description: str,
    ) -> None:
        """Test scoring falls back when embeddings fail."""
        with patch("backend.resume.scorer._get_embedding", new_callable=AsyncMock) as mock_embed:
            mock_embed.side_effect = Exception("API error")

            result = await score_resume(parsed_resume, sample_job_description)

            assert 0 <= result.total_score <= 100
            assert result.warnings  # Should have a warning about fallback

    @pytest.mark.asyncio
    async def test_score_resume_no_openai_key(
        self,
        parsed_resume: ParsedResume,
        sample_job_description: str,
    ) -> None:
        """Test scoring works without OpenAI key (fallback mode)."""
        with patch("backend.core.config.get_settings") as mock_settings:
            mock_settings.return_value.OPENAI_API_KEY = ""

            result = await score_resume(parsed_resume, sample_job_description)

            assert 0 <= result.total_score <= 100
            assert result.breakdown is not None


class TestParseResume:
    """Test resume parsing with real file formats."""

    @pytest.mark.asyncio
    async def test_parse_minimal_text_as_pdf(self) -> None:
        """Test that minimal text extraction is handled."""
        # Note: This tests the validation layer; actual PDF parsing requires a real PDF
        from backend.resume.parser import _validate_file
        ext = _validate_file("test.pdf", b"%PDF test content")
        assert ext == "pdf"

    def test_year_pattern(self) -> None:
        """Test year extraction pattern."""
        from backend.resume.parser import YEAR_PATTERN
        assert YEAR_PATTERN.search("Graduated 2016").group(0) == "2016"
        assert YEAR_PATTERN.search("Started 2022").group(0) == "2022"
